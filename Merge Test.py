
import os
import docx
from docx.shared import Inches
screendir = './Screenshots/'
testdir = './Test Cases/'
screenshots = [f for f in os.listdir(screendir) if '.png' in f]
appname = 'LO'
print(screenshots)
#for line in files:
#    fileno += 1
#    print(str(fileno) +". " + line)
#number = input('Please input the file number for the user stories file: ')
#userstories = files[int(number)-1]
for i in screenshots:
    #Get the App No.
    print(i[:2])
    appendixno=int(i[:2]) 
    #Get the Test ID
    idtest=i[2:14]
    doc = docx.Document(testdir+appname+ ' Test Results Appendix - Feature '+str(appendixno)+'.docx')
    for j in range(1,15):
        testrow = doc.tables[0].rows[j].cells[0]
        screenrow = doc.tables[0].rows[j].cells[2]
        print (testrow.text, idtest)
        if testrow.text == idtest:
            screenrow.add_paragraph().add_run().add_picture(
                    screendir+i,width=Inches(5),height=Inches(4.08))
            print(i +' was added')
            break
        
    doc.save(screendir+appname+ ' Test Results Appendix - Feature '+str(appendixno)+'.docx')
    '''for row in doc.tables[0].rows[1:]:
        for cell in row.cells:
            
            for paragraph in cell.paragraphs:
                print (paragraph.text)'''
'''    
doc = docx.Document(appname+ ' Test Results Appendix - Feature '+str(appendixno)+'.docx')
for i in range(1,len(featurescreenshot)+1):
    table = doc.tables[0].rows[i].cells[2].add_paragraph()

    table.add_run().add_picture(featurescreenshot[i-1],width=Inches(4), 
                  height=Inches(4.08))
#row.cells[2].text = doc.add_picture(screenshot)
    doc.save('test.docx')
    '''